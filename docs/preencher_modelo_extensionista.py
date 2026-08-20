from docx import Document
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph


TEMPLATE_PATH = r"d:\Users\remin\PetConectaB\docs\Atividades Extensionistas - Modelo de Proposta de tema e Trabalho Final.docx"
OUTPUT_PATH = r"d:\Users\remin\PetConectaB\docs\Atividades Extensionistas - Modelo de Proposta de tema e Trabalho Final - PREENCHIDO.docx"


def insert_paragraph_after(paragraph, text="", style=None):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if style is not None:
        new_para.style = style
    if text:
        new_para.add_run(text)
    return new_para


def find_paragraph(doc, startswith_text):
    for p in doc.paragraphs:
        txt = (p.text or "").strip()
        if txt.startswith(startswith_text):
            return p
    return None


def add_metodologia(doc):
    anchor = find_paragraph(doc, "Fonte:")
    if anchor is None:
        return

    textos = [
        "Metodologia aplicada: abordagem incremental e iterativa, com ciclos de planejamento, implementação, validação e refinamento.",
        "Justificativa: o projeto PetConecta já estava em evolução desde a Atividade Extensionista I, portanto a condução por ciclos curtos foi a forma mais adequada para consolidar melhorias sem interromper o desenvolvimento existente.",
        "Setor de aplicação: comunidade local (Coroa Vermelha - BA), com potencial de expansão para cidades vizinhas e organizações de proteção animal.",
        "Objetivo do gerenciamento: organizar o desenvolvimento do protótipo funcional, garantir rastreabilidade das decisões técnicas e manter alinhamento entre escopo, implementação e documentação acadêmica.",
        "Dados trabalhados: informações de pets (nome, status, localização, descrição, contato), registros de avistamentos e dados de autenticação de usuários.",
        "Recursos utilizados: HTML, CSS, JavaScript, Firebase (Authentication, Firestore, Storage, Hosting), GitHub e documentação técnica em arquivos de apoio.",
        "Passo a passo do projeto:",
        "1) Diagnóstico do problema local e definição do escopo.",
        "2) Estruturação da arquitetura e modelagem do sistema.",
        "3) Implementação incremental das páginas e funcionalidades principais.",
        "4) Integração com Firebase para autenticação, persistência e imagens.",
        "5) Validação funcional e ajustes de usabilidade, segurança e desempenho.",
        "6) Consolidação da documentação final e evidências de evolução.",
        "Diagrama textual da metodologia (fluxo):",
        "Levantamento do problema -> Planejamento do ciclo -> Implementação -> Validação -> Ajustes -> Documentação -> Entrega final",
    ]

    current = anchor
    for t in textos:
        current = insert_paragraph_after(current, t)


def add_resultados(doc):
    anchor = find_paragraph(doc, "Os resultados devem apresentar")
    if anchor is None:
        return

    textos = [
        "Resultados obtidos com a aplicação do projeto:",
        "1) Estruturação e disponibilização de um site navegável para divulgação de pets desaparecidos, encontrados e para adoção.",
        "2) Implementação de autenticação de usuários e área de gerenciamento de anúncios próprios.",
        "3) Integração com Firestore e Storage para registro de dados e upload de imagens.",
        "4) Organização da arquitetura e da modelagem técnica para manutenção e evolução do sistema.",
        "5) Melhoria contínua da solução com ajustes de segurança, privacidade de contato e desempenho de listagens.",
        "Evidências de comprovação:",
        "- Link do repositório GitHub: [INSERIR_LINK_DO_GITHUB_AQUI]",
        "- Link do vídeo (até 5 minutos): [INSERIR_LINK_DO_VIDEO_AQUI]",
        "- Documentos complementares no projeto: README.md, docs/modelagem-site.md, docs/arquitetura.md, docs/riscos-e-mitigacoes.md",
    ]

    current = anchor
    for t in textos:
        current = insert_paragraph_after(current, t)


def add_consideracoes(doc):
    anchor = find_paragraph(doc, "As considerações finais devem trazer")
    if anchor is None:
        return

    textos = [
        "Principais aprendizados e dificuldades do projeto:",
        "1) Aprendizado técnico: integração entre frontend e serviços Firebase para autenticação, banco de dados e armazenamento de arquivos.",
        "2) Aprendizado de projeto: importância do desenvolvimento incremental para corrigir e evoluir funcionalidades sem perder continuidade.",
        "3) Aprendizado de documentação: necessidade de manter modelagem, arquitetura e riscos atualizados para comprovar planejamento e execução.",
        "4) Dificuldade encontrada: alinhar evolução prática do sistema com formalização metodológica exigida na atividade.",
        "5) Superação da dificuldade: consolidação dos artefatos e registro das etapas de evolução para garantir coerência acadêmica.",
        "Conclusão: o projeto atendeu ao objetivo extensionista ao propor uma solução digital com potencial de impacto social, ao mesmo tempo em que fortaleceu competências técnicas e de gestão de projeto.",
    ]

    current = anchor
    for t in textos:
        current = insert_paragraph_after(current, t)


def main():
    doc = Document(TEMPLATE_PATH)
    add_metodologia(doc)
    add_resultados(doc)
    add_consideracoes(doc)
    doc.save(OUTPUT_PATH)
    print(OUTPUT_PATH)


if __name__ == "__main__":
    main()
