from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


def add_title(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(16)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_heading(doc, text, level=1):
    doc.add_heading(text, level=level)


def add_paragraph(doc, text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(10)


def main():
    doc = Document()

    add_title(doc, "PetConecta")
    add_paragraph(doc, "Seção acadêmica: Metodologia de Gerenciamento do Projeto e Evidências de Evolução")

    add_heading(doc, "1. Metodologia de Gerenciamento do Projeto", level=1)

    add_paragraph(
        doc,
        "Para o gerenciamento do projeto PetConecta, foi adotada uma abordagem incremental e iterativa, "
        "alinhada ao contexto acadêmico da Atividade Extensionista e ao histórico de desenvolvimento "
        "iniciado na etapa anterior. Essa metodologia foi escolhida por permitir evolução contínua da solução, "
        "com entregas progressivas, ajustes frequentes e melhoria da qualidade técnica e funcional ao longo do tempo.",
    )

    add_paragraph(
        doc,
        "Diferentemente de um modelo linear rígido, a abordagem incremental e iterativa possibilita construir "
        "o sistema em partes, validar resultados parciais e incorporar correções a cada ciclo. No caso do "
        "PetConecta, essa estratégia foi adequada porque o projeto exigiu refinamentos sucessivos de escopo, "
        "interface, segurança e organização dos dados, sem interromper o fluxo de desenvolvimento já em andamento "
        "desde a Extensionista I.",
    )

    add_paragraph(doc, "A aplicação prática da metodologia ocorreu em ciclos curtos, compostos por quatro etapas principais:")

    etapas = [
        "1) Planejamento do ciclo: definição dos objetivos imediatos, priorização de funcionalidades e identificação de riscos.",
        "2) Implementação: desenvolvimento ou ajuste de funcionalidades específicas, com foco em uma parte do sistema por vez.",
        "3) Validação: verificação funcional e técnica da entrega parcial, incluindo testes manuais de uso e análise de consistência dos dados.",
        "4) Refinamento: correção de falhas, melhoria de desempenho, revisão de segurança e atualização da documentação.",
    ]
    for e in etapas:
        add_paragraph(doc, e)

    add_paragraph(
        doc,
        "Esse formato permitiu, por exemplo, evoluir a solução em aspectos relevantes como centralização da "
        "configuração do Firebase, melhoria da consulta e renderização de dados, ajustes de autenticação, "
        "reforço das regras de segurança e proteção de informações sensíveis de contato. Assim, o projeto "
        "não ficou limitado a uma entrega única, mas consolidou um processo de amadurecimento contínuo.",
    )

    add_paragraph(
        doc,
        "Do ponto de vista de gerenciamento, a metodologia também favoreceu a rastreabilidade das decisões "
        "e da evolução do projeto. Os artefatos produzidos ao longo do desenvolvimento funcionam como "
        "evidências de planejamento e controle, incluindo modelagem do sistema, documento de arquitetura, "
        "análise de riscos e mitigação, histórico de alterações no repositório e organização das funcionalidades "
        "por prioridade e necessidade prática.",
    )

    add_paragraph(
        doc,
        "Em relação ao uso de frameworks específicos, destaca-se que o Kanban foi considerado como referência "
        "possível, mas não obrigatório. A escolha pela abordagem incremental e iterativa atende ao requisito "
        "da atividade por ser compatível com o tipo de projeto desenvolvido e com a dinâmica real de evolução adotada.",
    )

    add_paragraph(
        doc,
        "Conclui-se que a metodologia incremental e iterativa foi adequada ao PetConecta por permitir desenvolvimento "
        "progressivo, correções orientadas por validação contínua e fortalecimento da documentação do projeto. "
        "Essa escolha garantiu organização do trabalho, aderência aos objetivos da atividade extensionista e "
        "sustentação técnica para futuras evoluções do sistema.",
    )

    add_heading(doc, "2. Evidências de Evolução do Projeto", level=1)

    add_paragraph(
        doc,
        "A evolução do PetConecta foi conduzida de forma incremental e iterativa, com entregas parciais e "
        "refinamentos sucessivos ao longo do desenvolvimento. Para demonstrar o gerenciamento do projeto, "
        "esta seção apresenta os principais ciclos de evolução, os problemas identificados em cada etapa, "
        "as ações executadas e os resultados obtidos.",
    )

    add_paragraph(
        doc,
        "As evidências documentais que sustentam essa evolução estão registradas em README.md, "
        "docs/modelagem-site.md, docs/arquitetura.md e docs/riscos-e-mitigacoes.md, além do histórico de versões no GitHub.",
    )

    table = doc.add_table(rows=1, cols=7)
    table.style = "Table Grid"
    headers = [
        "Ciclo",
        "Período",
        "Objetivo do Ciclo",
        "Problema identificado",
        "Ação executada",
        "Evidência documental",
        "Resultado",
    ]

    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h

    rows = [
        [
            "1",
            "Extensionista I",
            "Estruturar a base do sistema",
            "Necessidade de centralizar a proposta da solução",
            "Definição do escopo inicial e estrutura do site",
            "Modelagem inicial e primeira versão no GitHub",
            "Base funcional do projeto estabelecida",
        ],
        [
            "2",
            "Início da Extensionista II",
            "Consolidar arquitetura técnica",
            "Organização técnica dispersa entre funcionalidades",
            "Revisão da arquitetura e padronização dos componentes",
            "docs/arquitetura.md",
            "Maior clareza estrutural e manutenção facilitada",
        ],
        [
            "3",
            "Desenvolvimento intermediário",
            "Melhorar persistência e consultas",
            "Crescimento de dados com risco de perda de desempenho",
            "Ajuste de consultas com ordenação e limite de leitura",
            "README.md",
            "Listagens mais estáveis e melhor desempenho",
        ],
        [
            "4",
            "Desenvolvimento intermediário",
            "Reforçar autenticação e fluxos de usuário",
            "Inconsistências entre primeiro acesso e retorno",
            "Ajuste no fluxo de autenticação e navegação",
            "README.md",
            "Experiência de acesso mais previsível",
        ],
        [
            "5",
            "Desenvolvimento avançado",
            "Reduzir exposição de dados sensíveis",
            "Contato público com risco de uso indevido",
            "Proteção de contato com fluxo de confirmação",
            "README.md e docs/arquitetura.md",
            "Maior privacidade sem perder funcionalidade",
        ],
        [
            "6",
            "Fase de consolidação",
            "Mapear riscos e respostas do projeto",
            "Necessidade de formalização de riscos",
            "Registro de riscos, causas e mitigação",
            "docs/riscos-e-mitigacoes.md",
            "Gestão de riscos documentada",
        ],
        [
            "7",
            "Entrega final",
            "Consolidar documentação acadêmica",
            "Necessidade de evidenciar método e evolução",
            "Atualização da modelagem por engenharia reversa e alinhamento metodológico",
            "docs/modelagem-site.md",
            "Documentação coerente com o sistema implementado",
        ],
    ]

    for row_data in rows:
        row_cells = table.add_row().cells
        for i, value in enumerate(row_data):
            row_cells[i].text = value

    add_heading(doc, "3. Análise da Evolução", level=1)

    add_paragraph(
        doc,
        "Os ciclos apresentados evidenciam que o projeto não foi desenvolvido em etapa única, mas sim em "
        "progressão contínua, com decisões orientadas por validação prática e necessidade real de ajuste. "
        "Esse comportamento é aderente à metodologia incremental e iterativa adotada no gerenciamento, "
        "pois cada ciclo agregou valor funcional e técnico ao sistema.",
    )

    add_paragraph(
        doc,
        "Além disso, a documentação acumulada ao longo do processo demonstra rastreabilidade: cada melhoria "
        "possui motivação, ação correspondente e resultado observável. Isso atende ao objetivo acadêmico de "
        "comprovar planejamento, execução, controle e evolução do projeto.",
    )

    add_heading(doc, "4. Conclusão", level=1)

    add_paragraph(
        doc,
        "As evidências de evolução confirmam que o PetConecta foi gerenciado com método, mesmo sem adoção "
        "obrigatória de um framework único como Kanban. A organização por ciclos, com documentação técnica "
        "e histórico de versão, comprova a estruturação do trabalho e a maturidade progressiva da solução desenvolvida.",
    )

    add_paragraph(
        doc,
        "Como material complementar de validação das evidências apresentadas, disponibiliza-se o repositório "
        "do projeto no GitHub, contendo histórico de alterações, versões e evolução do código-fonte ao longo "
        "dos ciclos de desenvolvimento.",
    )

    output_path = r"d:\\Users\\remin\\PetConectaB\\docs\\metodologia_e_evidencias_petconecta.docx"
    doc.save(output_path)
    print(output_path)


if __name__ == "__main__":
    main()
