from docx import Document

path = r"d:\Users\remin\PetConectaB\docs\Atividades Extensionistas - Modelo de Proposta de tema e Trabalho Final.docx"
doc = Document(path)

print("PARAGRAFOS:")
for i, p in enumerate(doc.paragraphs, 1):
    text = (p.text or "").strip()
    if text:
        print(f"{i}. {text}")

print("\nTABELAS:", len(doc.tables))
for ti, table in enumerate(doc.tables, 1):
    print(f"\nTabela {ti}: {len(table.rows)}x{len(table.columns)}")
    max_rows = min(20, len(table.rows))
    for r in range(max_rows):
        vals = []
        for c in table.rows[r].cells:
            vals.append((c.text or "").strip().replace("\n", " | "))
        print(" - " + " || ".join(vals))
